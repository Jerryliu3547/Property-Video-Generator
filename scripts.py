
import os
from google import genai
from docx import Document

# 1. Configuration
API_KEY = "AIzaSyAFUeHSOoLR5TuzGz5QDZAEiBRf6St1sOo" # Replace with your actual key
client = genai.Client(api_key=API_KEY, http_options={'api_version': 'v1'})

# 2. MLS DATA
mls_listing_text = """

Sample MlSlisting Description
""" # Trimmed for example, but keep your full text here

def generate_content(mls_content):
    # Prompt 1: The Voiceover Script
    vo_prompt = f"""
    You are a luxury real estate social media expert. 
    Create a 2-minute English voiceover script (approx 200 words) from this MLS data. 
    Style: Clear, warm, lifestyle-focused. No brackets or stage directions.
    Include contact line: "Please contact Sharon at 408-858-7703."
    MLS DATA: {mls_content}
    """

    print("Generating English Voiceover...")
    vo_response = client.models.generate_content(model="gemini-2.5-flash", contents=vo_prompt)
    vo_script = vo_response.text

    # Prompt 2: The Rednote (XHS) Post
    xhs_prompt = f"""
    You are a luxury real estate expert specializing in Rednote (Xiaohongshu).
    Based on the script below, create a viral Rednote Title and Description in Chinese.
    Use English for the street name and city. Include relevant emojis.
    SCRIPT: {vo_script}
    """

    print("Generating Chinese Rednote Post...")
    xhs_response = client.models.generate_content(model="gemini-2.5-flash", contents=xhs_prompt)
    xhs_post = xhs_response.text

    return vo_script, xhs_post

def save_to_single_docx(vo_content, xhs_content, filename="Property_Marketing_Package"):
    folder = "scripts"
    if not os.path.exists(folder):
        os.makedirs(folder)
    
    doc = Document()
    doc.add_heading("Real Estate Marketing Content", 0)

    # Section 1: English Script
    doc.add_heading("English Voiceover Script (TTS Ready)", level=1)
    doc.add_paragraph(vo_content)

    doc.add_page_break()

    # Section 2: Rednote Post
    doc.add_heading("Rednote (Xiaohongshu) Social Media Post", level=1)
    doc.add_paragraph(xhs_content)

    save_path = f"{folder}/{filename}.docx"
    doc.save(save_path)
    print(f"Success! Content saved to: {save_path}")

# --- Run Process ---
if __name__ == "__main__":
    # Generate both pieces of content
    english_vo, chinese_post = generate_content(mls_listing_text)
    
    # Save both into one single Word document
    save_to_single_docx(english_vo, chinese_post, "scripts")
